"""DOCX extractor using python-docx."""

import io
from typing import Optional
from uuid import UUID

from docx import Document

from src.services.extractors.base import DocumentExtractor, ExtractedChunk, ExtractedDocument


class DOCXExtractor(DocumentExtractor):
    """Extracts text and structure from DOCX files."""

    def supported_extensions(self) -> list[str]:
        return [".docx"]

    async def extract(
        self, file_bytes: bytes, file_key: str, document_id: UUID, version_id: UUID
    ) -> ExtractedDocument:
        doc = Document(io.BytesIO(file_bytes))

        text_parts = []
        chunks = []
        current_index = 0

        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                text_parts.append(para_text)

                heading_level = self._get_heading_level(para)
                chunks.append(
                    ExtractedChunk(
                        content=para_text,
                        chunk_index=current_index,
                        chunk_type="parent" if heading_level is None else "child",
                        section_title=para_text if heading_level else None,
                        section_level=heading_level,
                    )
                )
                current_index += 1

        table_texts = self._extract_tables(doc)
        full_text = "\n".join(text_parts + table_texts)
        language = self._detect_language(full_text)

        page_count = len(doc.sections) if doc.sections else 1

        return ExtractedDocument(
            text=full_text,
            chunks=chunks,
            page_count=page_count,
            language=language,
            metadata={"extractor": "python-docx", "paragraph_count": len(doc.paragraphs)},
        )

    def _extract_tables(self, doc: Document) -> list[str]:
        """Extract text from tables in the document."""
        table_texts = []

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(cells)
                table_texts.append(row_text)

        return table_texts

    def _get_heading_level(self, para) -> Optional[int]:
        """Detect heading level from paragraph style."""
        style_name = para.style.name.lower()

        if "heading 1" in style_name or style_name == "heading":
            return 1
        elif "heading 2" in style_name:
            return 2
        elif "heading 3" in style_name:
            return 3
        elif "heading 4" in style_name:
            return 4
        elif "heading 5" in style_name:
            return 5

        text = para.text.strip()
        if text.startswith("#"):
            if text.startswith("##"):
                return 2
            elif text.startswith("###"):
                return 3
            elif text.startswith("####"):
                return 4
            return 1

        return None

    def _detect_language(self, text: str) -> str:
        """Language detection based on character set - supports EN and VI."""
        if not text:
            return "en"
        vietnamese_chars = set("àáảãạăằắẳẵặâầấẩẫậèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡọùúủũụưừứửữựỳýỷỹỵđ")
        vietnamese_count = sum(1 for c in text.lower() if c in vietnamese_chars)
        return "vi" if vietnamese_count > len(text) * 0.05 else "en"